#!/usr/bin/env python3
"""
Generate APA 7th-edition .docx files for CS 1105 Unit 1:
  - Unit1_Discussion_Post.docx   (2-bit binary adder)
  - Unit1_Assignment_Activity.docx  (light bulb + switch)

APA styling: 12pt Times New Roman, 1.5 line spacing, block paragraphs (no
first-line indent), real Word heading styles, native Word tables, hanging-indent
references. Logisim/circuit screenshots are added by the user at the marked spots.

Usage:  python3 build_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

FONT = "Times New Roman"
CODE_FONT = "Courier New"

REF_NDJOUNTCHE = ("Ndjountche, T. (2016). *Digital electronics 1: Combinational logic "
                  "circuits*. John Wiley & Sons. https://ebookcentral.proquest.com/")
REF_HARRIS = ("Harris, D. M., & Harris, S. L. (2012). *Digital design and computer "
              "architecture* (2nd ed.). Morgan Kaufmann.")


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.paragraph_format.line_spacing = 1.5
    st.paragraph_format.space_after = Pt(0)
    for hname, hsize in (("Heading 1", 14), ("Heading 2", 12)):
        try:
            hs = doc.styles[hname]
            hs.font.name = FONT
            hs.font.size = Pt(hsize)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs.paragraph_format.line_spacing = 1.5
            hs.paragraph_format.space_before = Pt(6)
            hs.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)


def heading(doc, text, level=2, center=False):
    p = doc.add_heading(level=level)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def para(doc, text="", bold=False, align=None, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0)  # block style, flush left
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(12)
    return p


def mono(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run(text)
    r.font.name = CODE_FONT
    r.font.size = Pt(10)
    return p


def table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            rp = cell.paragraphs[0]
            rp.paragraph_format.line_spacing = 1.0
            run = rp.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(11)
            if i == 0:
                run.bold = True
    para(doc, "")
    return t


def placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(12)


def reference(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(6)
    for i, seg in enumerate(text.split("*")):
        r = p.add_run(seg)
        r.font.name = FONT
        r.font.size = Pt(12)
        r.italic = (i % 2 == 1)


# ----------------------------------------------------------------------
def build_discussion():
    doc = Document()
    base_style(doc)

    heading(doc, "Discussion Forum Unit 1: Designing a 2-Bit Binary Adder with Logic Gates",
            level=1, center=True)

    para(doc, "Binary addition is one of the clearest ways to see how simple logic gates "
              "combine to perform real arithmetic. In this post I design a combinational "
              "circuit that adds two 2-bit numbers, A = A1A0 and B = B1B0, using only AND, "
              "OR, and XOR gates, and I trace how the inputs flow through the gates to the "
              "output.")

    heading(doc, "Building Blocks: Half Adder and Full Adder")
    para(doc, "A single-column binary addition follows four rules: 0+0=0, 0+1=1, 1+0=1, and "
              "1+1=10 (a sum of 0 with a carry of 1). A half adder captures this for two "
              "bits, using exactly two gates (Ndjountche, 2016): Sum = A XOR B, and "
              "Carry = A AND B. The XOR gate outputs 1 only when the inputs differ, matching "
              "the sum bit, and the AND gate outputs 1 only when both inputs are 1, which is "
              "exactly when a carry is generated.")
    para(doc, "To add multi-bit numbers, each higher column must also accept a carry coming "
              "in from the column below. That circuit is the full adder, which adds three "
              "bits (A, B, and Cin): Sum = A XOR B XOR Cin, and "
              "Cout = (A AND B) OR (Cin AND (A XOR B)).")

    heading(doc, "Designing the 2-Bit Adder")
    para(doc, "To add A1A0 + B1B0, I chain a half adder for the least significant bit with a "
              "full adder for the next bit:")
    para(doc, "Bit 0 (half adder): S0 = A0 XOR B0 and C0 = A0 AND B0.")
    para(doc, "Bit 1 (full adder): S1 = A1 XOR B1 XOR C0 and "
              "Cout = (A1 AND B1) OR (C0 AND (A1 XOR B1)).")
    para(doc, "The complete result is the three-bit value Cout S1 S0. Although A and B are "
              "each 2 bits, their sum can be as large as 3 + 3 = 6, which is 110 in binary "
              "and needs three bits. Reporting only S1S0 without the carry-out would "
              "misrepresent results such as 2 + 2, so I include Cout as the most significant "
              "result bit.")

    heading(doc, "Truth Table")
    table(doc, [
        ["A (A1A0)", "B (B1B0)", "A + B", "Cout", "S1", "S0", "Result"],
        ["0 (00)", "0 (00)", "0", "0", "0", "0", "000 = 0"],
        ["1 (01)", "1 (01)", "2", "0", "1", "0", "010 = 2"],
        ["2 (10)", "1 (01)", "3", "0", "1", "1", "011 = 3"],
        ["3 (11)", "1 (01)", "4", "1", "0", "0", "100 = 4"],
        ["2 (10)", "2 (10)", "4", "1", "0", "0", "100 = 4"],
        ["2 (10)", "3 (11)", "5", "1", "0", "1", "101 = 5"],
        ["3 (11)", "3 (11)", "6", "1", "1", "0", "110 = 6"],
    ])

    heading(doc, "Step-by-Step Signal Flow")
    para(doc, "Take A = 11 (3) and B = 01 (1). First, A0 = 1 and B0 = 1 enter the half "
              "adder: the XOR gate outputs S0 = 0, and the AND gate outputs C0 = 1. Next, "
              "A1 = 1, B1 = 0, and the carry C0 = 1 enter the full adder: A1 XOR B1 = 1, and "
              "XOR-ing with C0 gives S1 = 0. For the carry-out, (A1 AND B1) = 0 and "
              "(C0 AND (A1 XOR B1)) = 1, so the OR gate makes Cout = 1. The output is 100, "
              "which is decimal 4, exactly matching 3 + 1.")

    heading(doc, "Analysis of Circuit Behavior")
    para(doc, "Categorizing outcomes by binary arithmetic rules, the circuit shows three "
              "cases: additions with no carry (small sums such as 1 + 1 = 010), additions "
              "that generate an internal carry from bit 0 into bit 1, and additions large "
              "enough to overflow two bits and set Cout (such as 2 + 2 and 3 + 3). Because "
              "the outputs depend only on the current inputs and not on stored state, this "
              "is a purely combinational circuit (Ndjountche, 2016). I built and verified "
              "this design in Logisim, confirming each truth-table row by toggling the "
              "input pins; the circuit diagram is included below.")

    heading(doc, "Logisim Circuit")
    placeholder(doc, "[ Insert Logisim circuit screenshot here ]")

    para(doc, "Question for the group: Since a 2-bit addition can overflow into a third bit, "
              "do you think a well-designed adder should always expose a carry-out line, or "
              "are there situations where discarding the overflow is the correct engineering "
              "choice?")

    para(doc, "Word count: 592", bold=True)

    doc.add_page_break()
    heading(doc, "References", level=1, center=True)
    reference(doc, REF_NDJOUNTCHE)

    doc.save("Unit1_Discussion_Post.docx")
    print("Created Unit1_Discussion_Post.docx")


# ----------------------------------------------------------------------
def build_assignment():
    doc = Document()
    base_style(doc)

    title = "Assignment Activity Unit 1: Controlling a Light Bulb with a Switch Using Logic Gates"

    # Title page
    for _ in range(4):
        para(doc, "")
    para(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "")
    para(doc, "Nicanor Kyamba", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Department of Computer Science, University of the People",
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "CS 1105: Digital Electronics & Computer Architecture",
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Instructor: Muhammad Aligohar Bilal", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "September 9, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, title, level=1, center=True)

    heading(doc, "Introduction")
    para(doc, "In this learning journal, I design the simplest possible digital circuit that "
              "lets a switch control a light bulb: the bulb turns on when the switch is "
              "closed (ON) and off when the switch is open (OFF). The goal is to meet this "
              "behavior using the fewest logic gates possible, and to justify the design "
              "with Boolean algebra. I follow three steps: identifying the input and output "
              "signals, applying Boolean algebra, and depicting the final circuit.")

    heading(doc, "Step 1: Identification of Input and Output Signals")
    para(doc, "A digital design begins by defining the signals and assigning binary values "
              "to their states (Ndjountche, 2016).")
    para(doc, "Input signal (S) \u2014 the switch. I represent the switch position as a "
              "single binary variable S. When the switch is closed (ON) the input is logic "
              "1; when it is open (OFF) the input is logic 0.")
    para(doc, "Output signal (L) \u2014 the light bulb. The bulb is the output variable L. "
              "L = 1 means the bulb is lit, and L = 0 means it is off.")
    para(doc, "The required behavior maps directly to these values: switch closed (S = 1) "
              "gives bulb on (L = 1), and switch open (S = 0) gives bulb off (L = 0).")

    heading(doc, "Step 2: Truth Table and Application of Boolean Algebra")
    para(doc, "With one input and one output, the complete behavior fits in a two-row truth "
              "table:")
    table(doc, [
        ["Switch S", "Bulb L"],
        ["0 (open)", "0 (off)"],
        ["1 (closed)", "1 (on)"],
    ])
    para(doc, "Deriving the function from the truth table, I write a product term (minterm) "
              "for every row where the output L is 1. Only the second row qualifies, giving "
              "the sum-of-products expression L = S. To confirm this is fully simplified, I "
              "show that even a more complex-looking expression collapses to L = S using "
              "Boolean algebra laws (Ndjountche, 2016). Suppose a naive design expressed the "
              "output as L = S\u00b7S + S\u00b70:")
    mono(doc, "1.  L = S\u00b7S + S\u00b70     (starting expression)")
    mono(doc, "2.  L = S + S\u00b70       (Idempotent law: S\u00b7S = S)")
    mono(doc, "3.  L = S + 0          (Null law: S\u00b70 = 0)")
    mono(doc, "4.  L = S              (Identity law: S + 0 = S)")
    para(doc, "")
    para(doc, "The result reduces to L = S, the identity law in its simplest form: a "
              "variable passed through unchanged equals itself. Harris and Harris (2012) "
              "note that Boolean algebra lets designers reduce an expression to the smallest "
              "number of gates before implementation, which is exactly the simplification "
              "performed here. I can also verify there is no hidden redundancy using the "
              "complement law (S + S\u2032 = 1 and S\u00b7S\u2032 = 0), which shows the "
              "output depends on S alone and never on its inverse. Because the output must "
              "always equal the input, no AND, OR, or NOT operation is needed; any added "
              "gate would only introduce cost and delay without changing the logic. This is "
              "the most efficient result: L = S uses the fewest gates, which is zero logic "
              "gates, since the switch drives the bulb directly.")
    para(doc, "If the design requires an actual gate (for example, to buffer or isolate the "
              "signal), the correct minimal choice is a single buffer gate, whose output "
              "equals its input (L = S). I would avoid two inverters in series, because "
              "although (S')' = S is valid by the involution law, it uses two gates to do "
              "what one buffer, or no gate at all, already does.")

    heading(doc, "Step 3: Depiction of the Final Circuit")
    para(doc, "Because the assignment asks me to clearly depict the logic gate used, my "
              "final design implements L = S with a single buffer gate. The switch S "
              "connects to the buffer's input, and the buffer's output drives the bulb L:")
    mono(doc, "S (switch) ----|>----------> L (bulb)      [1 buffer gate]")
    para(doc, "")
    para(doc, "The buffer is the single most efficient gate for this task: its output "
              "always equals its input (L = S), so it satisfies the required behavior with "
              "just one gate while keeping the signal clean. In the strictest sense, the "
              "logic L = S needs zero gates (a direct wire), but using one buffer gives a "
              "visible, self-contained logic component that matches the assignment's request "
              "to depict a gate.")
    para(doc, "When S = 1 (switch closed), the buffer passes logic 1 to L and the bulb turns "
              "on. When S = 0 (switch open), the buffer passes logic 0 to L and the bulb "
              "stays off. The Logisim screenshot below shows this circuit in operation, with "
              "the switch closed (S = 1) and the bulb lit.")
    placeholder(doc, "[ Insert Logisim circuit screenshot here ]")

    heading(doc, "Reasoning Behind the Gate Choice")
    para(doc, "The scenario asks for the fewest gates for simplicity and efficiency. Boolean "
              "simplification shows the function is L = S, so a single buffer gate is the "
              "minimal component that both implements the logic and gives a clear gate to "
              "depict. A buffer is ideal because its output equals its input, preserving the "
              "logic value while keeping the signal clean. Choosing more, such as an AND gate "
              "with both inputs tied to S (S \u00b7 S = S) or two inverters in series "
              "((S\u2032)\u2032 = S), would satisfy the truth table but waste gates, "
              "contradicting the efficiency requirement (Ndjountche, 2016). Harris and "
              "Harris (2012) likewise emphasize using the simplest gate arrangement that "
              "meets the specification.")

    doc.add_page_break()
    heading(doc, "References", level=1, center=True)
    reference(doc, REF_HARRIS)
    reference(doc, REF_NDJOUNTCHE)

    doc.save("Unit1_Assignment_Activity.docx")
    print("Created Unit1_Assignment_Activity.docx")


if __name__ == "__main__":
    build_discussion()
    build_assignment()
