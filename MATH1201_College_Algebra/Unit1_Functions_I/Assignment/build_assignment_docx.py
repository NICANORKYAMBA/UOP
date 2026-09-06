#!/usr/bin/env python3
"""
Generate the APA 7 .docx for MATH 1201 Unit 1 Written Assignment (Functions - I).

Meets the assignment's format requirements: Word document, double-spaced,
Times New Roman, 12pt, 1-inch margins. Includes an APA student title page,
real Word heading styles, and a placeholder for the GeoGebra graph (Task 2).

Usage:  python3 build_assignment_docx.py
Output: Unit1_Assignment_Activity.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

FONT = "Times New Roman"
MONO = "Courier New"
OUTPUT = "Unit1_Assignment_Activity.docx"

TITLE = "Written Assignment Unit 1: Functions - I"
AUTHOR = "Nicanor Kyamba"
AFFIL = "Department of Mathematics, University of the People"
COURSE = "MATH 1201: College Algebra"
INSTRUCTOR = "Instructor: Chibuike Agu"
DUE = "September 9, 2026"

REFS = [
    "Abramson, J. (2023). *Algebra and trigonometry* (2nd ed.). OpenStax. "
    "https://openstax.org/details/books/algebra-and-trigonometry-2e",
    "GeoGebra. (n.d.). *GeoGebra graphing calculator*. https://www.geogebra.org/calculator",
    "Stitz, C., & Zeager, J. (2013). *College algebra*. Stitz Zeager Open Source "
    "Mathematics. https://stitz-zeager.com/szca07042013.pdf",
    "Yoshiwara, K. (2020). *Modeling, functions, and graphs*. American Institute of "
    "Mathematics. https://yoshiwarabooks.org/mfg/colophon-1.html",
]


def base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.paragraph_format.line_spacing = 2.0  # APA double-spacing
    st.paragraph_format.space_after = Pt(0)
    for h, sz in (("Heading 1", 12), ("Heading 2", 12)):
        try:
            s = doc.styles[h]
            s.font.name = FONT
            s.font.size = Pt(sz)
            s.font.bold = True
            s.font.color.rgb = RGBColor(0, 0, 0)
            s.paragraph_format.line_spacing = 2.0
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(0)
        except KeyError:
            pass
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)


def para(doc, text="", bold=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(12)
    return p


def head(doc, text, level=2, center=False):
    # Uses the built-in Word "Heading 1"/"Heading 2" styles so the document
    # outline / Navigation pane recognizes them as native headings.
    p = doc.add_heading(level=level)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def title_style(doc, text):
    # Uses the built-in Word "Title" style (native), centered per APA.
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def mono(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(11)
    return p


def placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(12)


def reference(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    for i, seg in enumerate(text.split("*")):
        r = p.add_run(seg)
        r.font.name = FONT
        r.font.size = Pt(12)
        r.italic = (i % 2 == 1)


def build():
    doc = Document()
    base(doc)

    # Title page (APA student title page, uses native Title style for the title)
    for _ in range(4):
        para(doc, "")
    title_style(doc, TITLE)
    para(doc, "")
    para(doc, AUTHOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, AFFIL, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, COURSE, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, INSTRUCTOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, DUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    head(doc, TITLE, level=1, center=True)

    # Task 1
    head(doc, "Task 1: Interpreting the Graph (Domain, Range, One-to-One)")
    para(doc, "The given graph is a smooth curve that rises from the bottom of the plane on "
              "the left, peaks near the point (0, 5), dips to a small valley near (2, 3), "
              "and then rises steeply upward off the top of the plane on the right. The "
              "arrows on both ends indicate the curve continues without bound.")
    para(doc, "(i) Domain and Range. The curve extends left and right without end (both "
              "ends carry arrows), so every x-value is covered: the domain is all real "
              "numbers, (-infinity, infinity). The curve falls to negative infinity on the "
              "lower left and rises to positive infinity on the upper right, passing through "
              "every height in between, so the range is also all real numbers, "
              "(-infinity, infinity).")
    para(doc, "(ii) Function and one-to-one. Any vertical line touches the curve exactly "
              "once, so each input x gives exactly one output y; therefore the graph is a "
              "function (Abramson, 2023). However, because the curve rises, dips slightly "
              "between the peak near (0, 5) and the valley near (2, 3), then rises again, a "
              "horizontal line such as y = 4 crosses the curve three times. Since at least "
              "one horizontal line meets the graph more than once, the function is not "
              "one-to-one, so it has no inverse over its full domain (Abramson, 2023).")

    # Task 2
    head(doc, "Task 2: Avocado Export Function E(P) = P - 10000, P >= 10000")
    para(doc, "(i) Graph. E(P) = P - 10000 is a straight line of slope 1 with E-intercept "
              "-10000; since P >= 10000, the graph is the ray starting at (10000, 0) and "
              "rising to the right. It was plotted in GeoGebra using a scale of 1 unit = "
              "1000 on both axes.")
    placeholder(doc, "[ Insert GeoGebra graph of E(P) here ]")
    para(doc, "(ii) Is E(P) a function of P? Yes. Each production value P gives exactly one "
              "export value E(P) = P - 10000, and the line passes the vertical line test, "
              "so E is a function of P (Abramson, 2023).")
    para(doc, "(iii) Domain and range. Domain: P >= 10000, i.e., [10000, infinity). Range: "
              "at P = 10000, E = 0, and E grows without bound, so [0, infinity).")
    para(doc, "(iv) Export for 70 and 20 thousand of production. With values in thousands, "
              "P = 70000 gives E = 70000 - 10000 = 60000 (60 thousand units), and P = 20000 "
              "gives E = 20000 - 10000 = 10000 (10 thousand units).")
    para(doc, "(v) Variables. The independent variable is P (production), which is chosen "
              "freely; the dependent variable is E (export), whose value depends on P.")

    # Task 3
    head(doc, "Task 3: Rate of Change - Weights and Lengths of Two Animals")
    para(doc, "The graph shows a parabola (f) and a line (g) intersecting at A(5, 25), with "
              "x = length (feet) and y = weight (tons). A line through the origin and "
              "A(5, 25) is g(x) = 5x, and a parabola through the origin and A(5, 25) is "
              "f(x) = x squared.")
    para(doc, "(i) Rate of change at the intersection. The rate of change of weight with "
              "respect to length is the slope, slope = change in y over change in x. For "
              "the line g the rate of change is constant at 5 tons per foot everywhere, "
              "including at A. For the parabola f the rate of change is not constant; it "
              "increases as length increases. So animal g gains weight at a steady rate per "
              "foot, while animal f gains weight faster and faster as it lengthens.")
    para(doc, "(ii) Slopes of CD (on f) and EF (on g). Using slope = (y2 - y1)/(x2 - x1): "
              "on the parabola f(x) = x squared, take C = (2, 4) and D = (4, 16), giving "
              "slope CD = (16 - 4)/(4 - 2) = 12/2 = 6. On the line g(x) = 5x, take "
              "E = (1, 5) and F = (3, 15), giving slope EF = (15 - 5)/(3 - 1) = 10/2 = 5. "
              "The slope of EF is 5, matching the line's constant rate, so any two points on "
              "g give the same value; animal g gains weight at a steady 5 tons per foot. The "
              "slope of CD is 6, but this is only the average rate between x = 2 and x = 4 "
              "(between x = 1 and x = 3 it would be 4), confirming the parabola's rate of "
              "change is not constant. Animal f gains weight increasingly quickly as its "
              "length grows, rising much faster than g for larger lengths.")

    # Task 4
    head(doc, "Task 4: Local Extrema and Behavior of the Function")
    para(doc, "The graph is a smooth, repeating wave (a cosine-type curve) oscillating "
              "between a height of y = 1 at its peaks and y = -1 at its valleys. The labeled "
              "points are A(-11, 1), B(-8, -1), C(-5, 1), D(-2, -1), E(1.57, 1), "
              "F(4.71, -1), G(7.85, 1), and H(11, -1).")
    para(doc, "A local maximum is a point higher than all nearby points (a peak). Here the "
              "peaks A, C, E, and G are local maxima, each at y = 1. A local minimum is a "
              "point lower than all nearby points (a valley); here B, D, F, and H are local "
              "minima, each at y = -1 (Abramson, 2023). These differ from the absolute "
              "maximum and minimum, which are the single highest and lowest values over the "
              "entire domain. Because this curve repeats, every peak reaches the same height "
              "(1) and every valley the same depth (-1), so each peak ties for the absolute "
              "maximum and each valley ties for the absolute minimum. The distinction is "
              "that a local extremum need only be highest or lowest in its immediate "
              "neighborhood, while an absolute extremum is highest or lowest over the whole "
              "graph; a periodic function like this has infinitely many local extrema.")
    para(doc, "Intervals (reading left to right between the labeled points): increasing on "
              "(B, C) = (-8, -5); decreasing on (C, D) = (-5, -2); increasing on (D, E) = "
              "(-2, 1.57); decreasing on (E, F) = (1.57, 4.71); increasing on (F, G) = "
              "(4.71, 7.85); decreasing on (G, H) = (7.85, 11). At every peak the graph "
              "changes from increasing to decreasing (a local maximum), and at every valley "
              "from decreasing to increasing (a local minimum). The pattern continues "
              "indefinitely because the function is periodic.")

    # Task 5
    head(doc, "Task 5: Piecewise Tax Function for Country W")
    para(doc, "(i) The tax rule as a piecewise function. Let x be income and T(x) the tax:")
    mono(doc, "          | 0.10x,                        0 <= x <= 2200")
    mono(doc, "T(x) =    | 220 + 0.185(x - 2200),     2200 <  x <= 8945")
    mono(doc, "          | 1467.825 + 0.30(x - 8945),    x > 8945")
    para(doc, "For 2200 < x <= 8945, the first $2200 is taxed at 10% ($220) and the amount "
              "above $2200 at 18.5%. For x > 8945, the first $2200 gives $220, the band from "
              "$2200 to $8945 (width $6745) taxed at 18.5% gives $1,247.825, so the fixed "
              "part is $1,467.825, plus 30% of income above $8945.")
    para(doc, "(ii) Sample tax in each slab:")
    para(doc, "Slab a - income $2,000 (<= 2200): T = 0.10 x 2000 = $200.00.")
    para(doc, "Slab b - income $5,000 (2200 < x <= 8945): T = 220 + 0.185 x (5000 - 2200) = "
              "220 + 518 = $738.00.")
    para(doc, "Slab c - income $10,000 (> 8945): T = 1467.825 + 0.30 x (10000 - 8945) = "
              "1467.825 + 316.5 = $1,784.325, approximately $1,784.33.")
    para(doc, "Each result taxes only the income within each rate band, which is how a "
              "progressive piecewise tax is applied (Stitz & Zeager, 2013).")

    # References
    doc.add_page_break()
    head(doc, "References", level=1, center=True)
    for r in REFS:
        reference(doc, r)

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
