#!/usr/bin/env python3
"""
Generate formatted DOCX files for CS1102 Unit 6 submissions.
Reads the markdown files directly and converts them to DOCX with:
- Times New Roman, 12pt
- 1.5 line spacing
- 1" margins
- Code blocks in Courier New, 10pt, single-spaced
- Tables with Table Grid style
- No bookmarks
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import os
import re

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def set_document_defaults(doc):
    """Set margins for the document."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_heading_no_bookmark(doc, text, level=1):
    """Add a heading with Times New Roman, no bookmark."""
    heading = doc.add_heading('', level=level)
    run = heading.add_run(text)
    run.font.name = 'Times New Roman'
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_after = Pt(6)
    return heading


def add_normal_paragraph(doc, text, bold=False, italic=False):
    """Add a normal paragraph with Times New Roman 12pt, 1.5 spacing."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    return para


def add_code_line(doc, text):
    """Add a single code line with Courier New, 10pt, single-spaced."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    return para


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        para.paragraph_format.line_spacing = 1.5

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            para.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()
    return table


def parse_md_to_docx(md_path, docx_path):
    """Parse a markdown file and convert to formatted DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    set_document_defaults(doc)

    i = 0
    in_code_block = False
    in_table = False
    table_headers = []
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block start/end
        if line.startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                doc.add_paragraph()  # spacing after code
            else:
                # Start code block
                in_code_block = True
                doc.add_paragraph()  # spacing before code
            i += 1
            continue

        # Inside code block
        if in_code_block:
            add_code_line(doc, line)
            i += 1
            continue

        # Table detection
        if line.startswith('|') and '|' in line[1:]:
            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Check if next line is separator (|---|---|)
            if (i + 1 < len(lines) and
                re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip())):
                # This is a header row
                if in_table:
                    # Flush previous table
                    add_table(doc, table_headers, table_rows)
                in_table = True
                table_headers = cells
                table_rows = []
                i += 2  # skip header + separator
                continue
            elif in_table:
                # Data row
                table_rows.append(cells)
                i += 1
                continue

        # If we were in a table and hit a non-table line, flush it
        if in_table:
            add_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        # Headings
        if line.startswith('# '):
            add_heading_no_bookmark(doc, line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            add_heading_no_bookmark(doc, line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            add_heading_no_bookmark(doc, line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('#### '):
            add_heading_no_bookmark(doc, line[5:].strip(), level=4)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Bold paragraph (like **Level 1 — ...**)
        bold_match = re.match(r'^\*\*(.+)\*\*$', line.strip())
        if bold_match:
            add_normal_paragraph(doc, bold_match.group(1), bold=True)
            i += 1
            continue

        # Italic placeholder (like *[INSERT SCREENSHOT...]*)
        italic_match = re.match(r'^\*(.+)\*$', line.strip())
        if italic_match:
            add_normal_paragraph(doc, italic_match.group(1), italic=True)
            i += 1
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Clean markdown formatting from text
            text = clean_inline_markdown(text)
            para = doc.add_paragraph()
            para.style = 'List Bullet'
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            para.paragraph_format.line_spacing = 1.5
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if num_match:
            text = clean_inline_markdown(num_match.group(2))
            para = doc.add_paragraph()
            para.style = 'List Number'
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            para.paragraph_format.line_spacing = 1.5
            i += 1
            continue

        # Normal paragraph — collect continuation lines
        para_text = clean_inline_markdown(line.strip())
        i += 1
        # Don't merge lines — each line in md is its own paragraph
        add_normal_paragraph(doc, para_text)

    # Flush any remaining table
    if in_table:
        add_table(doc, table_headers, table_rows)

    doc.save(docx_path)
    return docx_path


def clean_inline_markdown(text):
    """Remove inline markdown formatting but keep the text."""
    # Bold + italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Links [text](url) -> text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


if __name__ == '__main__':
    print("Generating Unit 6 DOCX files from markdown...")
    print("-" * 50)

    # Discussion
    disc_md = os.path.join(BASE_DIR, 'Discussion', 'Unit6_Discussion_Assignment.md')
    disc_docx = os.path.join(BASE_DIR, 'Discussion', 'Unit6_Discussion_Assignment.docx')
    parse_md_to_docx(disc_md, disc_docx)
    print(f"✓ Discussion: {disc_docx}")

    # Assignment
    assign_md = os.path.join(BASE_DIR, 'Assignment', 'Unit6_Programming_Assignment.md')
    assign_docx = os.path.join(BASE_DIR, 'Assignment', 'Unit6_Programming_Assignment.docx')
    parse_md_to_docx(assign_md, assign_docx)
    print(f"✓ Assignment: {assign_docx}")

    print("-" * 50)
    print("\nFormatting applied:")
    print("  • Font: Times New Roman, 12pt")
    print("  • Line spacing: 1.5")
    print("  • Margins: 1 inch all sides")
    print("  • Code: Courier New, 10pt, single-spaced")
    print("  • Tables: Table Grid style")
    print("  • No bookmarks")
    print("\nRemember to insert screenshots in the Assignment DOCX.")
