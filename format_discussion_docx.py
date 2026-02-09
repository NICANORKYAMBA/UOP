from docx import Document
from docx.shared import Pt, Inches

doc = Document('/home/nicanorkyamba/UOP/ENGL1102_English_Composition_2/Unit2_Research_Sources/Discussion/Unit2_Discussion_Assignment_FORMATTED.docx')

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

for paragraph in doc.paragraphs:
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(12)
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

for paragraph in doc.paragraphs:
    paragraph.paragraph_format.line_spacing = 1.5

doc.save('/home/nicanorkyamba/UOP/ENGL1102_English_Composition_2/Unit2_Research_Sources/Discussion/Unit2_Discussion_Assignment_FORMATTED.docx')
