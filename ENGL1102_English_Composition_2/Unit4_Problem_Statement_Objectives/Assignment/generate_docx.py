#!/usr/bin/env python3
"""
Generate a properly formatted DOCX for ENGL1102 Unit 4 Assignment.
- Times New Roman 12pt
- Double-spaced
- 1-inch margins
- APA format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# --- Page Setup: 1-inch margins ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_paragraph_format(paragraph, bold=False, alignment=None, font_size=12,
                          space_after=0, first_line_indent=None, italic=False):
    """Set standard formatting for a paragraph."""
    pf = paragraph.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    if alignment:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Ensure Times New Roman works on all systems
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = doc.element.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')

def add_text(text, bold=False, alignment=None, font_size=12, space_after=0,
             first_line_indent=None, italic=False):
    """Add a simple paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_paragraph_format(p, bold=bold, alignment=alignment, font_size=font_size,
                         space_after=space_after, first_line_indent=first_line_indent,
                         italic=italic)
    return p

def add_mixed_paragraph(parts, first_line_indent=0.5, space_after=0):
    """Add a paragraph with mixed formatting (bold, italic, normal segments).
    parts = list of (text, bold, italic) tuples
    """
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = doc.element.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    return p


# ============================================================
# TITLE PAGE (APA Style)
# ============================================================
# Add blank lines to center title on page
for _ in range(6):
    add_text('', alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_text('Cybersecurity Investment and Data Breach Prevention in Small Businesses:',
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_text('Problem Statement, Research Objectives, and Questions',
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Blank line
add_text('', alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_text('Nicanor Kyamba', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_text('Department of Computer Science, University of the People',
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_text('ENGL1102: English Composition 2', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_text('February 2026', alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Page break
doc.add_page_break()


# ============================================================
# PROBLEM STATEMENT
# ============================================================
add_text('Problem Statement', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Paragraph 1 - Context
add_mixed_paragraph([
    ('Cybersecurity threats targeting small businesses have escalated at an alarming rate, exposing critical vulnerabilities in organizations that form the backbone of the global economy. The Verizon (2023) Data Breach Investigations Report found that "74% of all breaches involved the human element, including social engineering attacks, errors, or misuse" (p. 8), confirming that technological defenses alone are insufficient when employees remain the primary attack vector. Compounding this risk, the Cybersecurity and Infrastructure Security Agency (CISA, 2024) documented that organizations with structured employee training programs experience a 70% reduction in successful phishing attacks compared to those lacking formal security education. Despite this compelling evidence, small businesses continue to face disproportionate exposure because existing cybersecurity guidance often fails to account for their unique operational constraints.', False, False)
], first_line_indent=0.5)

# Paragraph 2 - The Gap
add_mixed_paragraph([
    ('Current literature offers valuable but fragmented approaches to this challenge. The National Institute of Standards and Technology (NIST, 2018) advocates for "a holistic approach integrating people, processes, and technology" (p. 3) through its five-function framework, yet its comprehensive scope presents implementation barriers for resource-limited organizations. Conversely, the Small Business Administration (SBA, 2023) recommends that "small businesses prioritize three high-impact, cost-effective measures: employee training, regular software updates, and multi-factor authentication" (p. 5), providing accessible entry-level guidance. However, Eling and Schnell (2016) argue that effective cybersecurity investment requires quantitative risk assessment methodologies to ensure informed decision-making, noting that organizations employing data-driven evaluation demonstrate superior incident response and resource allocation. While these perspectives collectively establish the importance of cybersecurity investment, they reveal a significant gap: limited research addresses the long-term sustainability of training programs, the measurement of behavioral change persistence, and how small businesses can progressively mature from foundational security measures toward comprehensive frameworks over time.', False, False)
], first_line_indent=0.5)

# Paragraph 3 - Significance
add_mixed_paragraph([
    ('This gap carries substantial consequences. Small businesses account for 99.9% of all U.S. firms and employ approximately 46% of the private-sector workforce (SBA, 2023). Without scalable, evidence-based models that integrate sustained training reinforcement with measurable progress indicators, these organizations remain disproportionately vulnerable to cyberattacks resulting in financial devastation, reputational harm, and operational disruption. Addressing this gap through the development of a tiered cybersecurity maturity model would provide small businesses with a structured, progressive pathway from basic security practices to robust risk management, ultimately strengthening the cybersecurity resilience of the broader economic landscape.', False, False)
], first_line_indent=0.5)


# ============================================================
# RESEARCH OBJECTIVES
# ============================================================
add_text('Research Objectives', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_mixed_paragraph([
    ('The overarching goal of this research is to evaluate the sustainability of cybersecurity training investments in small businesses and develop a progressive implementation framework. The following specific objectives guide this study:', False, False)
], first_line_indent=0.5)

# Objective 1
add_mixed_paragraph([
    ('1. To ', False, False),
    ('analyze', True, False),
    (' cybersecurity training approaches recommended for small businesses in government and peer-reviewed literature published between 2016 and 2024, assessing documented evidence of their effectiveness in reducing security incidents within a 24-month implementation window.', False, False)
], first_line_indent=0.5)

# Objective 2
add_mixed_paragraph([
    ('2. To ', False, False),
    ('evaluate', True, False),
    (' the extent to which existing cybersecurity frameworks, specifically NIST (2018), CISA (2024), and SBA (2023), address the scalability needs and resource constraints of small businesses with fewer than 50 employees.', False, False)
], first_line_indent=0.5)

# Objective 3
add_mixed_paragraph([
    ('3. To ', False, False),
    ('identify', True, False),
    (' specific gaps in current literature regarding long-term training sustainability, behavioral change measurement, and progressive security maturity pathways for resource-constrained organizations.', False, False)
], first_line_indent=0.5)

# Objective 4
add_mixed_paragraph([
    ('4. To ', False, False),
    ('propose', True, False),
    (' a tiered cybersecurity maturity model integrating continuous training reinforcement with quantifiable milestones, enabling small businesses to advance systematically from foundational practices toward comprehensive security frameworks.', False, False)
], first_line_indent=0.5)


# ============================================================
# RESEARCH QUESTIONS
# ============================================================
add_text('Research Questions', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_mixed_paragraph([
    ('Each research question directly corresponds to a specific objective to ensure alignment and coherence throughout the study:', False, False)
], first_line_indent=0.5)

# Question 1
add_mixed_paragraph([
    ('1. What cybersecurity training approaches does current literature recommend for small businesses, and what empirical evidence supports their effectiveness in reducing security incidents over sustained implementation periods? ', False, False),
    ('(Addresses Objective 1)', False, True)
], first_line_indent=0.5)

# Question 2
add_mixed_paragraph([
    ('2. To what extent do leading cybersecurity frameworks (NIST, CISA, SBA) accommodate the resource limitations and scalability requirements of businesses with fewer than 50 employees? ', False, False),
    ('(Addresses Objective 2)', False, True)
], first_line_indent=0.5)

# Question 3
add_mixed_paragraph([
    ('3. What specific gaps exist in the scholarly literature regarding the long-term sustainability of cybersecurity training programs and the measurement of persistent behavioral change among small business employees? ', False, False),
    ('(Addresses Objective 3)', False, True)
], first_line_indent=0.5)

# Question 4
add_mixed_paragraph([
    ('4. What essential components should a tiered cybersecurity maturity model incorporate to enable progressive, measurable, and sustainable security implementation in resource-constrained small businesses? ', False, False),
    ('(Addresses Objective 4)', False, True)
], first_line_indent=0.5)


# ============================================================
# REFERENCES (APA 7th Edition)
# ============================================================
doc.add_page_break()
add_text('References', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Reference 1 - CISA
p = doc.add_paragraph()
run = p.add_run('Cybersecurity and Infrastructure Security Agency. (2024). ')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
run = p.add_run('Cybersecurity best practices')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.font.italic = True
run = p.add_run('. U.S. Department of Homeland Security. https://www.cisa.gov/cybersecurity-best-practices')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
pf = p.paragraph_format
pf.line_spacing = 2.0; pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.first_line_indent = Inches(-0.5)
p.paragraph_format.left_indent = Inches(0.5)

# Reference 2 - Eling & Schnell
p = doc.add_paragraph()
run = p.add_run('Eling, M., & Schnell, W. (2016). What do we know about cyber risk and cyber risk insurance? ')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
run = p.add_run('The Journal of Risk Finance, 17')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.font.italic = True
run = p.add_run('(5), 474–491. https://doi.org/10.1108/JRF-09-2016-0122')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
pf = p.paragraph_format
pf.line_spacing = 2.0; pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.first_line_indent = Inches(-0.5)
p.paragraph_format.left_indent = Inches(0.5)

# Reference 3 - NIST
p = doc.add_paragraph()
run = p.add_run('National Institute of Standards and Technology. (2018). ')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
run = p.add_run('Framework for improving critical infrastructure cybersecurity')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.font.italic = True
run = p.add_run(' (Version 1.1). U.S. Department of Commerce. https://doi.org/10.6028/NIST.CSWP.04162018')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
pf = p.paragraph_format
pf.line_spacing = 2.0; pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.first_line_indent = Inches(-0.5)
p.paragraph_format.left_indent = Inches(0.5)

# Reference 4 - SBA
p = doc.add_paragraph()
run = p.add_run('Small Business Administration. (2023). ')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
run = p.add_run('Cybersecurity for small businesses')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.font.italic = True
run = p.add_run('. U.S. Small Business Administration. https://www.sba.gov/business-guide/manage-your-business/cybersecurity')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
pf = p.paragraph_format
pf.line_spacing = 2.0; pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.first_line_indent = Inches(-0.5)
p.paragraph_format.left_indent = Inches(0.5)

# Reference 5 - Verizon
p = doc.add_paragraph()
run = p.add_run('Verizon. (2023). ')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
run = p.add_run('2023 data breach investigations report')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.font.italic = True
run = p.add_run('. Verizon Business. https://www.verizon.com/business/resources/reports/dbir/')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
pf = p.paragraph_format
pf.line_spacing = 2.0; pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.first_line_indent = Inches(-0.5)
p.paragraph_format.left_indent = Inches(0.5)


# ============================================================
# SAVE
# ============================================================
output_path = '/home/nicanorkyamba/UOP/ENGL1102_English_Composition_2/Unit4_Problem_Statement_Objectives/Assignment/Unit4_Written_Assignment.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
