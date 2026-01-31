#!/usr/bin/env python3
"""
Format the discussion assignment DOCX file with proper formatting:
- Times New Roman, 12pt font
- 1.5 line spacing
- 1" margins on all sides
- Code blocks in Courier New
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

def format_discussion(input_file, output_file):
    doc = Document(input_file)
    
    # Set up document margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Format all paragraphs
    for paragraph in doc.paragraphs:
        # Set font to Times New Roman, 12pt
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Set line spacing to 1.5
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.5
        
        # Check if this is a code block
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in ['>>>', 'print', 'import ', '<class', 'SyntaxError', 'File', '0.5', '1.5', "'3.12"]):
            # Format as code
            for run in paragraph.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            # Single spacing for code blocks
            paragraph_format.line_spacing = 1.0
            # Add left indent
            paragraph_format.left_indent = Inches(0.5)
    
    # Format tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    paragraph.paragraph_format.line_spacing = 1.5
    
    # Save the formatted document
    doc.save(output_file)
    print(f"✓ Discussion assignment formatted and saved as: {output_file}")
    print("\nFormatting applied:")
    print("  - Font: Times New Roman, 12pt")
    print("  - Line spacing: 1.5")
    print("  - Margins: 1 inch on all sides")
    print("  - Code blocks: Courier New, 10pt, single-spaced")

if __name__ == "__main__":
    input_file = "Unit1_Discussion_Assignment.docx"
    output_file = "Unit1_Discussion_Assignment_FORMATTED.docx"
    
    try:
        format_discussion(input_file, output_file)
    except Exception as e:
        print(f"Error: {e}")
