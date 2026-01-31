#!/usr/bin/env python3
"""
Format the assignment DOCX file with proper APA formatting:
- Times New Roman, 12pt font
- Double spacing (2.0)
- 1" margins on all sides
- Code blocks in Courier New
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

def format_assignment(input_file, output_file):
    doc = Document(input_file)
    
    # Set up document margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Format all paragraphs
    for paragraph in doc.paragraphs:
        # Set font to Times New Roman, 12pt
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Set line spacing to double (2.0)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        # Check if this is a code block (starts with specific patterns)
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in ['print(', 'age =', 'city =', 'import ', '#', 'result', 'if ', 'elif ', 'else:']):
            # Format as code
            for run in paragraph.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            # Single spacing for code blocks
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # Add light gray background
            paragraph_format.left_indent = Inches(0.5)
    
    # Format tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
    
    # Save the formatted document
    doc.save(output_file)
    print(f"✓ Document formatted and saved as: {output_file}")
    print("\nFormatting applied:")
    print("  - Font: Times New Roman, 12pt")
    print("  - Line spacing: Double (2.0)")
    print("  - Margins: 1 inch on all sides")
    print("  - Code blocks: Courier New, 10pt, single-spaced")

if __name__ == "__main__":
    input_file = "Unit1_Programming_Assignment_FINAL.docx"
    output_file = "Unit1_Programming_Assignment_FORMATTED.docx"
    
    try:
        format_assignment(input_file, output_file)
    except Exception as e:
        print(f"Error: {e}")
        print("\nNote: You may need to install python-docx:")
        print("  pip install python-docx")
