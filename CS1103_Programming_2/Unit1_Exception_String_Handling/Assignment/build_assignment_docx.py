#!/usr/bin/env python3
"""
Generate an APA 7th-edition formatted .docx for the CS 1103 Unit 1
Programming Assignment (Text Analysis Tool).

- APA student title page (own page)
- 12pt Times New Roman, double-spaced body, 1-inch margins
- Source code in Courier New, single-spaced, in a light box
- Clear screenshot placeholders to paste images into
- References on their own page with hanging indent

Usage:  python3 build_assignment_docx.py
Output: Unit1_Programming_Assignment.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CODE_FILE = "TextAnalyzer.java"
OUTPUT = "Unit1_Programming_Assignment.docx"

# ---- Title page fields (EDIT the instructor name before submitting) ----
TITLE = "Programming Assignment Unit 1: Text Analysis Tool"
AUTHOR = "Nicanor Kyamba"
AFFILIATION = "Department of Computer Science, University of the People"
COURSE = "CS 1103: Programming 2"
INSTRUCTOR = "Instructor: Dr. Abeena Azad"
DUE_DATE = "September 9, 2026"

FONT = "Times New Roman"
CODE_FONT = "Courier New"


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    # Make built-in Heading styles use our font/size and proper spacing
    for hname, hsize in (("Heading 1", 14), ("Heading 2", 12)):
        try:
            hs = doc.styles[hname]
            hs.font.name = FONT
            hs.font.size = Pt(hsize)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs.paragraph_format.line_spacing = 1.5
            hs.paragraph_format.space_before = Pt(6)
            hs.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass


def set_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)


def add_page_number(doc):
    """APA student paper: page number in the top-right header."""
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run.font.name = FONT
    run.font.size = Pt(12)


def para(doc, text="", bold=False, align=None, size=12, space_after=10):
    """Block paragraph: 1.5 spacing, NO first-line indent, clear gap after."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(size)
    return p


def heading(doc, text, level=1, center=False):
    """Real Word heading style so it shows up as a document header/outline."""
    p = doc.add_heading(level=level)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def code_line(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.left_indent = Inches(0.25)
    r = p.add_run(text if text else "")
    r.font.name = CODE_FONT
    r.font.size = Pt(9)
    return p


def shade(paragraph, color="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def hanging_reference(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(6)
    # Support simple *italic* segments
    parts = text.split("*")
    for i, seg in enumerate(parts):
        r = p.add_run(seg)
        r.font.name = FONT
        r.font.size = Pt(12)
        r.italic = (i % 2 == 1)
    return p


def build():
    doc = Document()
    set_base_style(doc)
    set_margins(doc)
    add_page_number(doc)

    # ---------- TITLE PAGE ----------
    for _ in range(4):
        para(doc, "")
    para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "")
    para(doc, AUTHOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, AFFILIATION, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, COURSE, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, INSTRUCTOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, DUE_DATE, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------- BODY ----------
    heading(doc, TITLE, level=1, center=True)

    heading(doc, "Scenario", level=2)
    para(doc,
         "Create a text analysis tool that performs various operations on a given "
         "text input to help users gain insights into the text data through character "
         "and word analysis.")

    heading(doc, "Overview", level=2)
    para(doc,
         "This program is a text analysis tool written in Java. It reads a paragraph "
         "of text from the user and reports the total character count, the total word "
         "count, the most common character, the frequency of a user-chosen character, "
         "the frequency of a user-chosen word, and the number of unique words. Character "
         "and word comparisons are case-insensitive, and surrounding punctuation is "
         "ignored during word analysis so that \u201cmat\u201d and \u201cmat.\u201d are "
         "treated as the same word. All user input is validated: the paragraph cannot be "
         "empty, the character entry must be exactly one non-blank character, and the word "
         "entry must be a single word. The main workflow is wrapped in a try/catch/finally "
         "block so that any unexpected runtime error is reported cleanly and the Scanner "
         "resource is always closed.")

    heading(doc, "Tasks Implemented", level=2)
    tasks = [
        "User Input \u2014 the program asks the user to enter a paragraph of text and stores it.",
        "Character Count \u2014 displays the total number of characters in the text.",
        "Word Count \u2014 displays the total number of words (words separated by spaces).",
        "Most Common Character \u2014 displays the most frequent character; ties are broken by selecting one of the tied characters.",
        "Character Frequency \u2014 asks the user for a character and displays how many times it occurs (case-insensitive, so 'a' and 'A' are the same).",
        "Word Frequency \u2014 asks the user for a word and displays how many times it occurs (case-insensitive).",
        "Unique Words \u2014 displays the number of unique words in the text (case-insensitive).",
    ]
    for i, t in enumerate(tasks, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{i}. {t}")
        r.font.name = FONT
        r.font.size = Pt(12)

    # ---------- SOURCE CODE ----------
    heading(doc, "Source Code", level=2)
    with open(CODE_FILE, "r") as f:
        for raw in f.read().split("\n"):
            cp = code_line(doc, raw.replace("\t", "    "))
            shade(cp)

    # ---------- SAMPLE OUTPUT ----------
    heading(doc, "Sample Run and Output", level=2)
    para(doc, "Test input:", bold=True)
    for line in ["The cat sat on the mat. The cat is happy",
                 "t          (character to analyze)",
                 "the        (word to analyze)"]:
        shade(code_line(doc, line))
    para(doc, "")
    para(doc, "Program output:", bold=True)
    output = [
        "Enter a paragraph or a lengthy text:",
        "",
        "----- Text Analysis Results -----",
        "Total characters: 40",
        "Total words: 10",
        "Most common character: 't'",
        "",
        "Enter a single character to count its frequency:",
        "Frequency of 't': 7",
        "",
        "Enter a single word to count its frequency:",
        "Frequency of \"the\": 3",
        "Number of unique words: 7",
        "",
        "Analysis complete. Input closed.",
    ]
    for line in output:
        shade(code_line(doc, line))

    # ---------- SCREENSHOTS ----------
    heading(doc, "Screenshots of Output", level=2)
    para(doc, "Screenshot 1 \u2014 Program running with sample input and full output:",
         bold=True)
    para(doc, "[ Insert Screenshot 1 here ]", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "")
    para(doc, "Screenshot 2 \u2014 Input validation (invalid entry rejected, then re-prompted):",
         bold=True)
    para(doc, "[ Insert Screenshot 2 here ]", align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------- REFERENCES ----------
    doc.add_page_break()
    heading(doc, "References", level=1, center=True)
    hanging_reference(doc,
        "Eck, D. J. (2022). *Introduction to programming using Java* (Version 9, "
        "JavaFX ed.). Hobart and William Smith Colleges. "
        "https://math.hws.edu/javanotes/")
    hanging_reference(doc,
        "Morelli, R., & Wade, R. (2017). *Java, Java, Java: Object-oriented problem "
        "solving* (3rd ed., Chapter 10: Exceptions\u2014When things go wrong). "
        "LibreTexts. https://eng.libretexts.org/Bookshelves/Computer_Science/"
        "Programming_Languages/Java_Java_Java_-_Object-Oriented_Programming_"
        "(Morelli_and_Walde)/10%3A_Exceptions-_When_Things_Go_Wrong")

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
