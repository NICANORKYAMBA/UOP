#!/usr/bin/env python3
"""
Generate an APA 7th-edition formatted .docx for the CS 1103 Unit 1
Discussion Forum post (Exception Handling in Java).

Forum posts do not use a title page. This produces a clean, APA-styled
document: 12pt Times New Roman, double-spaced, the post body with short
code illustrations in monospace, the word count, and a References section
with hanging indent.

Usage:  python3 build_discussion_docx.py
Output: Unit1_Discussion_Post.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUTPUT = "Unit1_Discussion_Post.docx"
FONT = "Times New Roman"
CODE_FONT = "Courier New"

TITLE = "Discussion Forum Unit 1: Exception Handling in Java"

# Body: list of ("text", type) where type is 'p' (paragraph) or 'code'
BODY = [
    ("When I mentor someone brand new to Java, I start with a promise: exceptions are not "
     "your enemy, they are Java\u2019s way of telling you the truth about what went wrong. "
     "An exception is an object that represents an error or an anomalous condition that "
     "appears while a program is running and interrupts its normal flow (Morelli & Wade, "
     "2017). Without a plan, that interruption crashes the program; with a plan, the program "
     "responds calmly. Eck (2022) frames the try..catch statement as exactly that plan, a "
     "way to catch and respond to errors instead of stopping. I explain the four keywords "
     "using one running scenario: a program that asks a user for their age.", "p"),

    ("try \u2014 attempt the risky part. I tell my mentee that try is where you say, "
     "\u201cI\u2019ll attempt this, knowing it might fail.\u201d The classic risky step is "
     "turning user text into a number, because the user might type letters. My everyday "
     "analogy is carrying a full cup of coffee across a room: you are hopeful, but you stay "
     "ready for a spill.", "p"),
    ("try {", "code"),
    ("    int age = Integer.parseInt(userInput);   // may throw", "code"),
    ("    System.out.println(\"Next year you'll be \" + (age + 1));", "code"),
    ("}", "code"),

    ("catch \u2014 the safety net. If the risky line fails, Java throws an exception and "
     "jumps to the matching catch, which is the towel you keep nearby for the spill.", "p"),
    ("catch (NumberFormatException e) {", "code"),
    ("    System.out.println(\"Please enter a whole number, digits only.\");", "code"),
    ("}", "code"),

    ("Now, instead of crashing, the program explains the problem. I also show beginners that "
     "exceptions are organized in a hierarchy under Throwable, so you can catch a very "
     "specific type like NumberFormatException or a broader parent like Exception "
     "(Morelli & Wade, 2017). I encourage catching the specific type, because a precise net "
     "gives a precise, helpful message.", "p"),

    ("throw \u2014 raise your own alarm. Sometimes the value is a valid number but still "
     "makes no sense, such as a negative age. Here I decide it is an error and signal it "
     "with throw. Morelli and Wade (2017) compare throwing an exception to pulling a fire "
     "alarm to announce an abnormal condition.", "p"),
    ("if (age < 0) {", "code"),
    ("    throw new IllegalArgumentException(\"Age cannot be negative.\");", "code"),
    ("}", "code"),

    ("finally \u2014 the cleanup crew. The finally block runs no matter what happens: "
     "success, a caught error, or an uncaught one on its way up. It is where you release "
     "resources like files, network connections, or a Scanner. My analogy is turning off "
     "the stove before leaving the kitchen, whether dinner was a triumph or a disaster.", "p"),
    ("Scanner keyboard = new Scanner(System.in);", "code"),
    ("try {", "code"),
    ("    // read and process input", "code"),
    ("} catch (Exception e) {", "code"),
    ("    System.out.println(\"Something went wrong: \" + e.getMessage());", "code"),
    ("} finally {", "code"),
    ("    keyboard.close();   // always runs", "code"),
    ("}", "code"),

    ("The design lesson I most want a beginner to absorb comes straight from the reading: "
     "robust programs plan for errors from the earliest stages of development rather than "
     "adding handling as an afterthought (Morelli & Wade, 2017). In practice that means "
     "validating input first so fewer exceptions occur, catching specific types, giving "
     "users clear feedback, and never leaving a catch block empty, because an empty handler "
     "hides bugs instead of fixing them. This same care carries into string work, where "
     "Java\u2019s immutable String methods and the mutable StringBuilder are used "
     "constantly, and choosing the right one keeps programs both correct and efficient "
     "(Samoylov, 2018).", "p"),

    ("Put simply: try attempts, catch recovers, throw raises, and finally cleans up. "
     "Together they turn an unavoidable reality, that things go wrong, into something a "
     "program can face gracefully. That difference is what separates fragile code from "
     "software people can actually trust.", "p"),

    ("My question for the group: When you catch an exception, how do you decide whether to "
     "handle it right where it happens or to declare throws and let the calling method deal "
     "with it, and does that decision change for checked versus unchecked exceptions?", "p"),
]

WORD_COUNT = "Word count: 615"

REFERENCES = [
    "Eck, D. J. (2022). *Introduction to programming using Java* (Version 9, JavaFX ed.). "
    "Hobart and William Smith Colleges. https://math.hws.edu/javanotes/",
    "Morelli, R., & Wade, R. (2017). *Java, Java, Java: Object-oriented problem solving* "
    "(3rd ed., Chapter 10: Exceptions\u2014When things go wrong). LibreTexts. "
    "https://eng.libretexts.org/Bookshelves/Computer_Science/Programming_Languages/"
    "Java_Java_Java_-_Object-Oriented_Programming_(Morelli_and_Walde)/"
    "10%3A_Exceptions-_When_Things_Go_Wrong",
    "Samoylov, N. (2018). *Introduction to programming: Learn to program in Java with data "
    "structures, algorithms, and logic*. Packt Publishing.",
]


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    # Make built-in Heading 1 use our font/size/color and 1.5 spacing
    try:
        h1 = doc.styles["Heading 1"]
        h1.font.name = FONT
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.line_spacing = 1.5
        h1.paragraph_format.space_before = Pt(6)
        h1.paragraph_format.space_after = Pt(6)
    except KeyError:
        pass

    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Title as a real Heading 1 (centered)
    ht = doc.add_heading(level=1)
    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ht.paragraph_format.line_spacing = 1.5
    r = ht.add_run(TITLE)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 0, 0)

    for text, kind in BODY:
        if kind == "p":
            pp = doc.add_paragraph()
            pp.paragraph_format.line_spacing = 1.5
            pp.paragraph_format.first_line_indent = Inches(0)
            pp.paragraph_format.space_after = Pt(6)
            rr = pp.add_run(text)
            rr.font.name = FONT
            rr.font.size = Pt(12)
        else:
            cp = doc.add_paragraph()
            cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.left_indent = Inches(0.5)
            rr = cp.add_run(text)
            rr.font.name = CODE_FONT
            rr.font.size = Pt(10)

    # Word count
    wc = doc.add_paragraph()
    wc.paragraph_format.line_spacing = 1.5
    wc.paragraph_format.space_before = Pt(6)
    r = wc.add_run(WORD_COUNT)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(12)

    # References on a new page (Heading 1, centered)
    doc.add_page_break()
    hp = doc.add_heading(level=1)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.line_spacing = 1.5
    r = hp.add_run("References")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 0, 0)

    for ref in REFERENCES:
        rp = doc.add_paragraph()
        pf = rp.paragraph_format
        pf.line_spacing = 1.5
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.space_after = Pt(6)
        for i, seg in enumerate(ref.split("*")):
            run = rp.add_run(seg)
            run.font.name = FONT
            run.font.size = Pt(12)
            run.italic = (i % 2 == 1)

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
